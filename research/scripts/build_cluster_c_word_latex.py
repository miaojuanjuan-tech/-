from pathlib import Path
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\Lenovo\Documents\Codex\2026-07-04\new-chat-2")
OUT = ROOT / "outputs"

TITLE = "从虚拟物品到数字自我：乙游角色资产、情绪价值与持续消费"
SUBTITLE = "Cluster C 论文写作支架与 LaTeX 排版源文件（scaffold only）"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_run(paragraph, text, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_east_asia(run, font="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font)
    return run


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    add_run(p, label + "：", bold=True, color="1F4D78")
    add_run(p, text)
    doc.add_paragraph()


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "内容槽"
    for c in hdr:
        set_cell_shading(c, "E8EEF5")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_width(table, [1.5, 4.8])
    doc.add_paragraph()


def read_csv(name):
    with open(OUT / name, newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        s = styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def build_docx():
    constructs = read_csv("cluster_c_construct_ledger.csv")
    items = read_csv("cluster_c_minimum_survey_items.csv")
    decisions = read_csv("cluster_c_design_decision_register.csv")

    doc = Document()
    setup_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    r.font.name = "Calibri"
    set_east_asia(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")
    set_east_asia(r)

    add_note(
        doc,
        "作者边界",
        "本文档是论文写作支架，不是最终论文正文。所有正式段落、结论措辞和投稿文本必须由作者自行撰写和确认。",
    )

    doc.add_heading("1. 研究定位", level=1)
    add_kv_table(
        doc,
        [
            ("主线", "Cluster C：数字自我、虚拟商品与游戏消费。"),
            ("工作题目", TITLE),
            ("核心路径", "虚拟角色资产的数字自我延伸 -> 虚拟物品情感价值 -> 角色依恋/品牌依恋 -> 持续消费。"),
            ("研究对象", "近 3 或 6 个月玩过国内女性向恋爱手游的成年玩家。"),
            ("分析单位", "首选玩家-最喜欢角色关系；备选为玩家层面。"),
            ("报告边界", "横截面问卷只能支持关联和机制一致性，不能写成因果效应。"),
        ],
    )

    doc.add_heading("2. 引言写作骨架", level=1)
    intro_rows = [
        ("P1 现象入口", "描述乙游作为女性向数字娱乐/情感消费场景。需要作者补行业事实、平台现象或玩家行为材料。"),
        ("P2 理论入口", "用体验消费、享乐价值、品牌体验说明消费不仅是功能性购买。引用 Holbrook & Hirschman 1982; Babin et al. 1994; Brakus et al. 2009。"),
        ("P3 Cluster C 入口", "把角色、卡牌、账号资产、剧情记忆连接到数字自我、虚拟物品和 loved objects。引用 Belk 1988; Belk 2013; Ahuvia 2005。"),
        ("P4 研究缺口", "现有文献分散在体验消费、品牌关系、虚拟商品、游戏消费；缺少乙游角色资产如何转化为情绪价值和持续消费的整合。"),
        ("P5 研究问题", "槽位：本研究考察 [乙游成年玩家] 中 [数字自我延伸] 如何与 [虚拟物品情感价值]、[角色/品牌依恋] 和 [持续消费] 相关。"),
        ("P6 方法概览", "槽位：质性探索/预测试 + 问卷 SEM/PLS-SEM + 可选文本分析。只写实际会做的数据。"),
        ("P7 贡献槽", "理论贡献：把数字自我和虚拟商品文献推进到乙游情感消费场景；方法贡献和实践贡献需等数据支持后再收紧。"),
    ]
    add_kv_table(doc, intro_rows)

    doc.add_heading("3. 文献综述结构", level=1)
    clusters = [
        ("3.1 体验消费与情绪价值", "主张槽：乙游消费可从体验消费和情绪价值理解。证据：Holbrook & Hirschman; Babin; Brakus。边界：这些文献不是乙游场景。"),
        ("3.2 数字自我与虚拟物品", "主张槽：角色、卡牌、账号资产可作为数字自我延伸或 loved objects。证据：Belk; Ahuvia; Hamari & Keronen; Lehdonvirta。边界：平台控制改变了所有权含义。"),
        ("3.3 角色依恋与品牌/IP依恋", "主张槽：玩家可能与角色和品牌/IP形成不同层次的关系。证据：Fournier; Park; Batra; Labrecque。边界：角色依恋不等于品牌依恋。"),
        ("3.4 社群与二创实践", "主张槽：社群参与和二创实践可能共同生产角色亲密感。证据：Kozinets; Zhou et al. 2024。边界：不能作为总体比例证据。"),
    ]
    add_kv_table(doc, clusters)

    doc.add_heading("4. 构念与变量表", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdrs = ["构念", "模型角色", "工作定义", "文献锚点", "风险/下一步"]
    for i, h in enumerate(hdrs):
        table.cell(0, i).text = h
        set_cell_shading(table.cell(0, i), "E8EEF5")
    for c in constructs:
        cells = table.add_row().cells
        cells[0].text = c["construct_name"]
        cells[1].text = c["role_in_model"]
        cells[2].text = c["working_definition"]
        cells[3].text = c["literature_anchor"]
        cells[4].text = f"{c['risk_note']}；{c['next_action']}"
    set_table_width(table, [1.0, 1.05, 2.1, 1.25, 1.1])
    doc.add_paragraph()

    doc.add_heading("5. 假设/研究命题槽位", level=1)
    add_kv_table(
        doc,
        [
            ("H1 槽", "数字自我延伸与虚拟物品情感价值之间的正向关联。"),
            ("H2 槽", "虚拟物品情感价值与角色依恋之间的正向关联。"),
            ("H3 槽", "角色依恋与持续消费意向之间的正向关联。"),
            ("H4 槽", "虚拟物品情感价值和角色依恋在数字自我延伸与持续消费之间起链式中介作用。"),
            ("边界变量槽", "商业化反感/操纵感可能削弱情绪价值到持续消费意向的路径。"),
            ("语言边界", "若只有横截面数据，全部写成 relationship/association，不写 cause/effect。"),
        ],
    )

    doc.add_heading("6. 最小问卷题项", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdrs = ["题项ID", "构念", "题项文本", "量表/备注"]
    for i, h in enumerate(hdrs):
        table.cell(0, i).text = h
        set_cell_shading(table.cell(0, i), "E8EEF5")
    for it in items:
        cells = table.add_row().cells
        cells[0].text = it["item_id"]
        cells[1].text = it["construct_name"]
        cells[2].text = it["item_text_cn"]
        cells[3].text = f"{it['response_scale']}；{it['notes']}"
    set_table_width(table, [0.6, 1.1, 3.55, 1.05])
    doc.add_paragraph()

    doc.add_heading("7. 方法与分析计划槽位", level=1)
    add_kv_table(
        doc,
        [
            ("预测试", "先做 50-80 人预测试，检查题项理解、缺失率、量表分布和初步判别效度。"),
            ("样本表", "报告年龄段、乙游玩龄、付费经历、主推角色、招募渠道。"),
            ("测量模型", "报告 Cronbach alpha、CR、AVE、HTMT、载荷。"),
            ("主模型", "SEM/PLS-SEM：数字自我延伸 -> 情绪价值 -> 角色依恋 -> 持续消费意向。"),
            ("替代模型", "数字自我延伸 -> 角色依恋 -> 情绪价值 -> 持续消费意向。"),
            ("稳健性", "区分付费/非付费玩家、高频/低频玩家；结果替换为自报消费行为。"),
        ],
    )

    doc.add_heading("8. 作者决策登记", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdrs = ["决策ID", "组件", "当前选择", "风险", "作者需决定"]
    for i, h in enumerate(hdrs):
        table.cell(0, i).text = h
        set_cell_shading(table.cell(0, i), "E8EEF5")
    for d in decisions:
        cells = table.add_row().cells
        cells[0].text = d["decision_id"]
        cells[1].text = d["design_component"]
        cells[2].text = d["current_choice"]
        cells[3].text = d["risk_if_wrong"]
        cells[4].text = d["author_decision_needed"]
    set_table_width(table, [0.6, 1.0, 1.8, 1.8, 1.1])
    doc.add_paragraph()

    doc.add_heading("9. 不支持或需谨慎的写法", level=1)
    cautions = [
        ("不要写", "数字自我延伸导致持续消费。"),
        ("可写槽", "数字自我延伸与持续消费意向呈关联，且路径关系与理论机制一致。"),
        ("不要写", "所有乙游玩家都会把角色视为自我延伸。"),
        ("可写槽", "在受访样本或访谈材料中，部分玩家将角色资产描述为身份、记忆或投入的延伸。"),
        ("不要写", "平台商业化一定构成情感操纵。"),
        ("可写槽", "情感商业化可能伴随商业化反感、付费压力和操纵感，需作为边界条件检验。"),
    ]
    add_kv_table(doc, cautions)

    doc.add_heading("10. LaTeX 排版说明", level=1)
    add_note(
        doc,
        "排版",
        "同目录已生成 XeLaTeX/ctexart 源文件。LaTeX 文件保留相同章节结构、表格槽和作者填写提示。",
    )

    path = OUT / "cluster_c_manuscript_scaffold.docx"
    doc.save(path)
    return path


def esc_tex(s):
    return (
        s.replace("\\", r"\textbackslash{}")
        .replace("&", r"\&")
        .replace("%", r"\%")
        .replace("$", r"\$")
        .replace("#", r"\#")
        .replace("_", r"\_")
        .replace("{", r"\{")
        .replace("}", r"\}")
    )


def build_tex():
    constructs = read_csv("cluster_c_construct_ledger.csv")
    items = read_csv("cluster_c_minimum_survey_items.csv")

    tex = []
    tex.append(r"\documentclass[UTF8,zihao=-4]{ctexart}")
    tex.append(r"\usepackage[a4paper,margin=2.5cm]{geometry}")
    tex.append(r"\usepackage{booktabs,longtable,array,tabularx,enumitem,hyperref,xcolor}")
    tex.append(r"\hypersetup{colorlinks=true,linkcolor=blue,citecolor=blue,urlcolor=blue}")
    tex.append(r"\setlist[itemize]{leftmargin=2em}")
    tex.append(r"\setlist[enumerate]{leftmargin=2em}")
    tex.append(r"\title{" + TITLE + r"\\{\large Cluster C 论文写作支架（非最终正文）}}")
    tex.append(r"\author{作者待填写}")
    tex.append(r"\date{\today}")
    tex.append(r"\begin{document}")
    tex.append(r"\maketitle")
    tex.append(r"\begin{center}\fbox{\parbox{0.9\textwidth}{\textbf{作者边界：} 本文档是写作支架，不是最终论文正文。所有正式段落、结论措辞和投稿文本必须由作者自行撰写和确认。}}\end{center}")
    tex.append(r"\begin{abstract}")
    tex.append(r"\noindent 摘要槽：作者在完成研究设计、数据收集和分析后填写。当前版本只保留研究对象、核心路径、方法和贡献的填写提示。")
    tex.append(r"\end{abstract}")
    tex.append(r"\section{研究定位}")
    tex.append(r"\begin{itemize}")
    tex.append(r"\item 主线：Cluster C，数字自我、虚拟商品与游戏消费。")
    tex.append(r"\item 核心路径：虚拟角色资产的数字自我延伸 $\rightarrow$ 虚拟物品情感价值 $\rightarrow$ 角色依恋/品牌依恋 $\rightarrow$ 持续消费。")
    tex.append(r"\item 研究对象：近3或6个月玩过国内女性向恋爱手游的成年玩家。")
    tex.append(r"\item 分析单位：首选玩家-最喜欢角色关系。")
    tex.append(r"\item 报告边界：横截面问卷只能支持关联和机制一致性，不能写成因果效应。")
    tex.append(r"\end{itemize}")
    tex.append(r"\section{引言写作骨架}")
    tex.append(r"\begin{enumerate}")
    for item in [
        "P1 现象入口：描述乙游作为女性向数字娱乐/情感消费场景；作者补行业事实或平台数据。",
        "P2 理论入口：体验消费、享乐价值和品牌体验说明消费包含幻想、感受和乐趣。",
        "P3 Cluster C 入口：角色、卡牌、账号资产和剧情记忆连接到数字自我和虚拟物品。",
        "P4 研究缺口：现有文献分散，缺少乙游角色资产如何转化为情绪价值和持续消费的整合。",
        "P5 研究问题：填写 population、exposure、mediator、outcome。",
        "P6 方法概览：只写实际会做的数据和方法。",
        "P7 贡献槽：等数据支持后再收紧贡献强度。",
    ]:
        tex.append(r"\item " + esc_tex(item))
    tex.append(r"\end{enumerate}")
    tex.append(r"\section{文献综述结构}")
    tex.append(r"\subsection{体验消费与情绪价值}")
    tex.append("主张槽：乙游消费可从体验消费和情绪价值理解。引用槽：Holbrook \\& Hirschman (1982); Babin et al. (1994); Brakus et al. (2009)。")
    tex.append(r"\subsection{数字自我与虚拟物品}")
    tex.append("主张槽：角色、卡牌、账号资产可作为数字自我延伸或 loved objects。引用槽：Belk (1988, 2013); Ahuvia (2005)。")
    tex.append(r"\subsection{角色依恋与品牌/IP依恋}")
    tex.append("主张槽：玩家可能与角色和品牌/IP形成不同层次的关系。引用槽：Fournier (1998); Park et al. (2010); Batra et al. (2012)。")
    tex.append(r"\subsection{社群与二创实践}")
    tex.append("主张槽：社群参与和二创实践可能共同生产角色亲密感。引用槽：Kozinets (2002); Zhou et al. (2024)。")
    tex.append(r"\section{构念与变量表}")
    tex.append(r"\begin{longtable}{p{2.5cm}p{2.2cm}p{4.8cm}p{3.2cm}p{2.8cm}}")
    tex.append(r"\toprule 构念 & 模型角色 & 工作定义 & 文献锚点 & 风险/下一步 \\ \midrule")
    tex.append(r"\endfirsthead \toprule 构念 & 模型角色 & 工作定义 & 文献锚点 & 风险/下一步 \\ \midrule \endhead")
    for c in constructs:
        tex.append(
            esc_tex(c["construct_name"]) + " & " +
            esc_tex(c["role_in_model"]) + " & " +
            esc_tex(c["working_definition"]) + " & " +
            esc_tex(c["literature_anchor"]) + " & " +
            esc_tex(c["risk_note"] + "；" + c["next_action"]) + r" \\"
        )
    tex.append(r"\bottomrule\end{longtable}")
    tex.append(r"\section{研究命题槽位}")
    tex.append(r"\begin{enumerate}")
    for h in [
        "H1 槽：数字自我延伸与虚拟物品情感价值之间的正向关联。",
        "H2 槽：虚拟物品情感价值与角色依恋之间的正向关联。",
        "H3 槽：角色依恋与持续消费意向之间的正向关联。",
        "H4 槽：虚拟物品情感价值和角色依恋的链式中介。",
        "边界变量槽：商业化反感/操纵感可能削弱情绪价值到持续消费意向的路径。",
    ]:
        tex.append(r"\item " + esc_tex(h))
    tex.append(r"\end{enumerate}")
    tex.append(r"\section{最小问卷题项}")
    tex.append(r"\begin{longtable}{p{1.4cm}p{2.5cm}p{7.5cm}p{3.2cm}}")
    tex.append(r"\toprule 题项ID & 构念 & 题项文本 & 量表/备注 \\ \midrule")
    tex.append(r"\endfirsthead \toprule 题项ID & 构念 & 题项文本 & 量表/备注 \\ \midrule \endhead")
    for it in items:
        tex.append(
            esc_tex(it["item_id"]) + " & " +
            esc_tex(it["construct_name"]) + " & " +
            esc_tex(it["item_text_cn"]) + " & " +
            esc_tex(it["response_scale"] + "；" + it["notes"]) + r" \\"
        )
    tex.append(r"\bottomrule\end{longtable}")
    tex.append(r"\section{分析计划槽位}")
    tex.append(r"\begin{itemize}")
    for a in [
        "预测试：50-80人，检查题项理解、缺失率和初步判别效度。",
        "测量模型：报告 Cronbach alpha、CR、AVE、HTMT、载荷。",
        "主模型：SEM/PLS-SEM 检验数字自我延伸 -> 情绪价值 -> 角色依恋 -> 持续消费意向。",
        "替代模型：数字自我延伸 -> 角色依恋 -> 情绪价值 -> 持续消费意向。",
        "稳健性：区分付费/非付费玩家、高频/低频玩家；结果替换为自报消费行为。",
    ]:
        tex.append(r"\item " + esc_tex(a))
    tex.append(r"\end{itemize}")
    tex.append(r"\section{不支持或需谨慎的写法}")
    tex.append(r"\begin{itemize}")
    for c in [
        "不要写：数字自我延伸导致持续消费。可写：数字自我延伸与持续消费意向呈关联。",
        "不要写：所有乙游玩家都会把角色视为自我延伸。可写：在样本或访谈材料中，部分玩家如此描述。",
        "不要写：平台商业化一定构成情感操纵。可写：情感商业化可能伴随商业化反感和付费压力。",
    ]:
        tex.append(r"\item " + esc_tex(c))
    tex.append(r"\end{itemize}")
    tex.append(r"\section{作者待决策事项}")
    tex.append(r"\begin{enumerate}")
    for d in [
        "分析单位：玩家，还是玩家-角色关系？",
        "品牌/IP依恋是否保留，还是先聚焦角色依恋？",
        "是否收集真实或截图验证消费行为？",
        "商业化反感/操纵感放入主模型，还是作为讨论边界？",
    ]:
        tex.append(r"\item " + esc_tex(d))
    tex.append(r"\end{enumerate}")
    tex.append(r"\end{document}")

    path = OUT / "cluster_c_manuscript_scaffold.tex"
    path.write_text("\n".join(tex), encoding="utf-8")
    return path


if __name__ == "__main__":
    docx_path = build_docx()
    tex_path = build_tex()
    print(docx_path)
    print(tex_path)
